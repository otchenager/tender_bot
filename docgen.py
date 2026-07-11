"""Tender document package generation — HUMAN IN THE LOOP ONLY.

IRON RULE (legal): nothing in this module submits anything anywhere.
BY government tenders require the director's digital signature and carry
company liability, so the generated package is downloaded by the USER,
reviewed, and submitted manually through the official portal. The
"Подать" button in the UI only records the user's own per-tender decision.

Package contents (ZIP):
  1. ценовое_предложение.xlsx — smeta positions priced with the
     contractor's OWN price list (styled after смет_аВлада.pdf).
  2. техническое_предложение.docx — AI-drafted text based on the tender,
     explicitly marked as a draft for review.
  3. заявка.docx — application form filled with company profile data.
"""

import io
import zipfile
from datetime import datetime

import db
from logger import get_logger

log = get_logger("docgen")


# ---------------------------------------------------------------------------
# 1. Price proposal (xlsx)
# ---------------------------------------------------------------------------

def _build_price_proposal(tender: dict, positions: list[dict], profile: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Ценовое предложение"

    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    bold = Font(bold=True)

    ws["A1"] = profile.get("company_name") or "«Наименование организации»"
    ws["A1"].font = Font(bold=True, size=13)
    ws["A2"] = f"УНП: {profile.get('unp') or '—'}   {profile.get('address') or ''}"
    ws["A4"] = "ЦЕНОВОЕ ПРЕДЛОЖЕНИЕ"
    ws["A4"].font = Font(bold=True, size=14)
    ws["A5"] = f"по тендеру: {tender.get('title') or ''}"
    ws["A6"] = f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}"

    headers = ["№", "Наименование работ/материалов", "Ед. изм.", "Кол-во",
               "Цена за ед., BYN", "Сумма, BYN", "Примечание"]
    header_row = 8
    for col, title in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=title)
        cell.font = bold
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    total = 0.0
    row = header_row + 1
    for i, p in enumerate(positions, start=1):
        qty = p.get("smeta_quantity") or 0
        my_price = p.get("matched_my_price")
        amount = round(my_price * qty, 2) if (my_price is not None and qty) else None
        if amount is not None:
            total += amount
        note = ""
        if my_price is None:
            note = "нет в прайсе — заполнить вручную"
        elif p.get("match_status") == "yellow":
            note = "проверить соответствие позиции"

        values = [i, p.get("smeta_name") or "", p.get("smeta_unit") or "",
                  qty, my_price, amount, note]
        for col, v in enumerate(values, start=1):
            cell = ws.cell(row=row, column=col, value=v)
            cell.border = border
            if col in (4, 5, 6):
                cell.number_format = "#,##0.00"
        row += 1

    ws.cell(row=row, column=2, value="ИТОГО:").font = bold
    total_cell = ws.cell(row=row, column=6, value=round(total, 2))
    total_cell.font = bold
    total_cell.number_format = "#,##0.00"

    ws.cell(row=row + 2, column=2,
            value=f"Директор ____________________ {profile.get('director') or ''}")

    widths = [5, 52, 9, 9, 15, 15, 30]
    for col, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=header_row, column=col).column_letter].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 2. Technical proposal (docx, AI-drafted)
# ---------------------------------------------------------------------------

_TECH_PROMPT = """\
Ты — инженер-сметчик белорусской строительной компании.
Напиши текст технического предложения для участия в тендере.

Тендер: {title}
Основные виды работ (из сметы):
{works}

Структура (4-6 абзацев, деловой стиль, русский язык):
1. Краткое представление подрядчика и готовность выполнить работы.
2. Понимание объёма и состава работ.
3. Технология и организация выполнения (последовательность, контроль качества).
4. Сроки и гарантийные обязательства (общими формулировками).
5. Соответствие ТНПА Республики Беларусь.

Пиши ТОЛЬКО текст предложения, без заголовка и подписи."""


def _build_technical_proposal(tender: dict, positions: list[dict], profile: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    works = "\n".join(
        f"- {p.get('smeta_name')} ({p.get('smeta_quantity') or ''} {p.get('smeta_unit') or ''})"
        for p in positions[:20]
    )

    body_text = None
    try:
        from file_processor import _call_claude
        body_text = _call_claude(_TECH_PROMPT.format(
            title=tender.get("title") or "", works=works or "—",
        ))
    except Exception as e:
        log.error(f"Technical proposal AI draft failed: {e}")
    if not body_text:
        body_text = (
            "Организация подтверждает готовность выполнить полный комплекс работ, "
            "предусмотренных тендерной документацией, в установленные сроки и в "
            "соответствии с требованиями ТНПА Республики Беларусь.\n\n"
            "[AI-черновик недоступен — дополните текст вручную.]"
        )

    doc = Document()
    h = doc.add_heading("Техническое предложение", level=1)
    doc.add_paragraph(f"по тендеру: {tender.get('title') or ''}")
    doc.add_paragraph(f"Претендент: {profile.get('company_name') or '—'}, "
                      f"УНП {profile.get('unp') or '—'}")
    doc.add_paragraph("")

    for para in body_text.split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("")
    warn = doc.add_paragraph()
    run = warn.add_run(
        "⚠ ЧЕРНОВИК, сгенерирован AI — обязательно проверьте и отредактируйте "
        "перед подачей. Удалите эту пометку."
    )
    run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph(f"Директор ____________________ {profile.get('director') or ''}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 3. Application form (docx)
# ---------------------------------------------------------------------------

def _build_application(tender: dict, profile: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Заявление на участие в процедуре закупки", level=1)
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Предмет закупки: {tender.get('title') or '—'}")
    if tender.get("url"):
        doc.add_paragraph(f"Ссылка на процедуру: {tender['url']}")
    doc.add_paragraph("")

    rows = [
        ("Полное наименование претендента", profile.get("company_name")),
        ("УНП", profile.get("unp")),
        ("Юридический адрес", profile.get("address")),
        ("Руководитель", profile.get("director")),
        ("Телефон", profile.get("phone")),
        ("E-mail", profile.get("email")),
        ("Банк", profile.get("bank_name")),
        ("Расчётный счёт (IBAN)", profile.get("bank_account")),
        ("БИК банка", profile.get("bank_code")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value or "—"

    doc.add_paragraph("")
    doc.add_paragraph(
        "Настоящим подтверждаем согласие с условиями проведения процедуры "
        "закупки и готовность выполнить работы в соответствии с требованиями "
        "документации."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Директор ____________________ {profile.get('director') or ''}")
    doc.add_paragraph("М.П.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_package(tender_id: int) -> tuple[bytes, str] | None:
    """Build the full ZIP package for a tender. Returns (zip_bytes, filename)
    or None when the tender does not exist."""
    tender = db.get_tender(tender_id)
    if not tender:
        return None
    positions = db.get_tender_positions(tender_id)
    profile = db.get_company_profile()

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("ценовое_предложение.xlsx",
                    _build_price_proposal(tender, positions, profile))
        zf.writestr("техническое_предложение.docx",
                    _build_technical_proposal(tender, positions, profile))
        zf.writestr("заявка.docx", _build_application(tender, profile))
        zf.writestr(
            "README.txt",
            "Пакет документов подготовлен автоматически Тендер-ботом.\n"
            "ОБЯЗАТЕЛЬНО проверьте каждый документ перед подачей.\n"
            "Подача выполняется вручную через официальный портал с ЭЦП директора.\n"
            f"Тендер: {tender.get('title') or ''}\n"
            f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n",
        )

    filename = f"tender_{tender_id}_documents.zip"
    log.info(f"Document package generated for tender {tender_id}")
    return zip_buf.getvalue(), filename
